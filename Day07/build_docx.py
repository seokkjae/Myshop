# -*- coding: utf-8 -*-
"""모닝브루 사업 소개서 DOCX 생성 (PRD-lite v0.1 + 시장조사 + 경쟁사 비교 기반)"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- 색상 토큰 (홈페이지 디자인과 통일: 따뜻한 커피 브라운) ----
ACCENT = RGBColor(0xA8, 0x5A, 0x28)      # 진한 브라운
ACCENT2 = RGBColor(0xC0, 0x70, 0x3A)     # 밝은 브라운
INK = RGBColor(0x3A, 0x31, 0x2A)
INK_SOFT = RGBColor(0x6B, 0x5D, 0x52)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEAD_FILL = "C0703A"
SUBHEAD_FILL = "F6E7D4"

FONT = "맑은 고딕"

doc = Document()

# 기본 폰트 (서양/한글 모두 맑은 고딕)
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def set_run_font(run, size=10.5, bold=False, color=INK, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def para(text="", size=10.5, bold=False, color=INK, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p

def heading(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    # 좌측 강조 바 느낌: 번호를 accent 색으로
    r1 = p.add_run(f"{num}. ")
    set_run_font(r1, size=15, bold=True, color=ACCENT2)
    r2 = p.add_run(text)
    set_run_font(r2, size=15, bold=True, color=ACCENT)
    # 하단 경계선
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "ECE2D6")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

def add_table(headers, rows, header_fill=HEAD_FILL, widths=None, first_col_bold=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=WHITE)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            bold = first_col_bold and i == 0
            set_run_font(r, size=9.5, bold=bold,
                         color=ACCENT if bold else INK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def add_toc_field():
    """실제 목차 필드 (Word에서 F9로 갱신). 갱신 전에는 placeholder 표시."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-1" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "[목차] Word에서 이 영역을 클릭 후 F9 키를 누르면 자동 생성됩니다."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(el)

# =========================================================
# 1. 표지
# =========================================================
for _ in range(4):
    doc.add_paragraph()
para("☕", size=48, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("모닝브루", size=40, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Morning Brew", size=14, color=INK_SOFT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("망원동에서 가장 일찍 켜지는 불빛 —", size=14, bold=True, color=ACCENT2,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("당신의 아침을 여는 한 잔.", size=14, bold=True, color=ACCENT2,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
para("사업 소개서", size=16, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para("서울 망원동(마포구) · 출근길 동네 카페", size=11, color=INK_SOFT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("작성일: 2026-06-20  ·  작성자: admin@inno-curve.com", size=10, color=INK_SOFT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("버전 v0.1 (Draft)", size=9, color=INK_SOFT, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# =========================================================
# 목차
# =========================================================
para("목   차", size=18, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
toc_items = [
    "1.  사업 개요",
    "2.  시장 현황",
    "3.  서비스 소개",
    "4.  연락처 및 다음 단계",
]
for it in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(1.5)
    r = p.add_run(it)
    set_run_font(r, size=12, color=INK)
para("", space_after=8)
add_toc_field()  # Word용 실제 자동목차 필드 (선택적으로 F9 갱신)

doc.add_page_break()

# =========================================================
# 1. 사업 개요 (문제 / 타깃 / 가치 제안 — 3행 표)
# =========================================================
heading(1, "사업 개요")
para("서울 망원동 골목 끝의 동네 카페 「모닝브루」는, 출근 전 동네를 지나는 직장인과 거주민에게 "
     "이른 아침 갓 내린 커피와 가벼운 아침거리를 빠르게 제공하고 단골 관계로 묶어내는 것을 목표로 합니다. "
     "현재 초기 컨셉/MVP 검증 단계입니다.", space_after=12)

add_table(
    ["구분", "내용"],
    [
        ["문제 (Problem)",
         "출근 전 짧은 시간(7~9시)에 동네에서 만족스러운 커피 경험을 얻기 어렵다. "
         "① 출근 피크의 대기 줄, ② 이른 아침 문 연 동네 카페 부족, ③ 빠른 픽업·테이크아웃 비최적화, "
         "④ 함께 먹을 아침거리 부족, ⑤ 단골을 알아보지 못함(적립·주문 저장 부재)."],
        ["타깃 (Target)",
         "1차: 망원동을 지나 출근하는 20~40대 직장인(시간 압박·테이크아웃 중심). "
         "2차: 이른 아침 산책·운동 후 들르는 동네 거주민. "
         "공통적으로 가격보다 시간·신선함·익숙함을 중시."],
        ["가치 제안 (Value)",
         "가격 경쟁이 아닌 비가격 축으로 승부 — 이른 아침 시간대 선점 + 동네 단골 관계 + "
         "아침 루틴 경험. 사전주문·빠른 픽업 동선으로 줄 서지 않고, 커피+간단 푸드 세트와 "
         "간편 적립으로 '나를 알아주는 동네 카페'를 제공."],
    ],
    widths=[3.5, 12.5],
    first_col_bold=True,
)

# =========================================================
# 2. 시장 현황 (시장조사 요약 + 경쟁사 비교표)
# =========================================================
heading(2, "시장 현황")
para("타깃 고객 불편(페인) 요약", size=12, bold=True, color=ACCENT2, space_before=4, space_after=6)
para("출근 전 짧은 시간이라는 공통 맥락에서 5가지 핵심 불편이 도출되었습니다. "
     "(고객 심리·니즈 서술은 일반적 이용 패턴에 기반한 추정이며, 가격·영업시간 등은 현장 검증이 필요합니다.)",
     size=9.5, color=INK_SOFT, space_after=8)

add_table(
    ["#", "고객 불편 (Pain)", "현재 대안", "모닝브루의 해법"],
    [
        ["1", "출근 시간대 대기 줄로 시간 부족", "프랜차이즈 앱 사이렌오더, 편의점 커피", "사전주문·픽업존 등 빠른 동선"],
        ["2", "이른 아침 문 연 동네 카페가 드묾", "일찍 여는 프랜차이즈·편의점, 홈카페", "경쟁사보다 이른 오픈 시간"],
        ["3", "빠른 테이크아웃에 최적화 안 됨", "큰길가 매장, 개인 텀블러 지참", "테이크아웃 전용 픽업 동선·홀더"],
        ["4", "함께 먹을 아침거리 부족", "베이커리·편의점 따로 들름, 결식", "커피+간단 푸드(베이글·스콘) 세트"],
        ["5", "단골이 되어도 매번 주문 설명", "프랜차이즈 앱, 종이 쿠폰", "'나의 주문' 저장·간편 적립"],
    ],
    widths=[0.8, 4.7, 5.0, 5.0],
)

para("", space_after=6)
para("경쟁사·대체재 비교", size=12, bold=True, color=ACCENT2, space_before=8, space_after=6)
para("모닝브루는 가격(저가 프랜차이즈·편의점)으로 이기기 어려우므로, 시간대 선점·단골 관계·아침 루틴 경험이라는 "
     "다른 축에서 경쟁합니다. (가격·리뷰 수 등 시점에 따라 변하는 값은 ‘확인 필요’.)",
     size=9.5, color=INK_SOFT, space_after=8)

add_table(
    ["경쟁 대안", "가격대(확인 필요)", "강점", "약점", "차별화 지점"],
    [
        ["스타벅스 (망원·합정)", "아메리카노 4,500원대", "브랜드 신뢰, 넓은 좌석, 사이렌오더", "개성 없음, 혼잡, 비쌈", "동네 단골을 알아보는 1:1 감성"],
        ["개인 감성 카페 (골목)", "5,000~6,500원", "인테리어·SNS 노출, 시그니처", "가격 높음, 늦은 오픈, 좌석 적음", "이른 아침 + 합리적 가격으로 빈 시간대 공략"],
        ["저가 프랜차이즈\n(메가·컴포즈·빽다방)", "1,500~2,500원", "압도적 저가, 테이크아웃 속도", "맛·경험 평준화, 머무는 공간 아님", "가격 대신 품질·공간·아침 루틴으로 경쟁"],
        ["중가 프랜차이즈\n(이디야·투썸)", "3,000~5,000원", "무난한 맛, 디저트 라인업", "차별점 약함, 충성도 낮음", "원두 스토리 있는 ‘콘텐츠 한 잔’"],
        ["편의점 커피 (CU·GS25)", "1,000~1,800원", "초저가, 즉시성, 어디에나", "맛·향 한계, 경험 0", "아침을 여는 의식(ritual)으로 포지셔닝"],
        ["배달 앱 (배민·쿠팡이츠)", "음료가+배달비 3,000원~", "안 나가도 됨, 묶음 주문", "배달비 부담, 식은 커피", "걸어서 픽업 + 갓 내린 신선함 + 단골 적립"],
    ],
    widths=[3.2, 2.6, 3.4, 3.0, 3.8],
    first_col_bold=True,
)

# =========================================================
# 3. 서비스 소개 (핵심 흐름 단계별)
# =========================================================
heading(3, "서비스 소개")
para("모닝브루의 핵심 고객 경험은 ‘출근 전 5분’ 안에 완결되도록 설계되었습니다. "
     "사전 주문부터 픽업, 단골화까지의 흐름은 다음과 같습니다.", space_after=12)

steps = [
    ("STEP 1 · 사전 주문", "출근길에 모바일/전화로 메뉴와 픽업 시간을 미리 지정합니다. 줄 서지 않는 경험의 출발점."),
    ("STEP 2 · 얼리 오픈 준비", "경쟁 동네 카페보다 이른 시간에 개점해, 갓 내린 커피와 당일 구운 아침거리를 준비합니다."),
    ("STEP 3 · 빠른 픽업", "테이크아웃 전용 픽업존과 간편 결제로 대기를 최소화. 골목 끝 매장이라도 동선을 짧게."),
    ("STEP 4 · 모닝 세트 제공", "갓 내린 커피 + 간단 푸드(베이글·스콘) 세트로 아침 결식 고객의 ‘두 번 들름’을 해결."),
    ("STEP 5 · 단골 인식·적립", "‘나의 주문’ 저장과 간편 적립으로 매번 설명할 필요 없이, 나를 알아주는 동네 카페 경험."),
    ("STEP 6 · 아침 루틴화", "원두 스토리·단골 이름 기억 등 비가격 콘텐츠로 재방문을 습관(ritual)으로 정착."),
]
for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run("● ")
    set_run_font(r1, size=10.5, bold=True, color=ACCENT2)
    r2 = p.add_run(title)
    set_run_font(r2, size=11, bold=True, color=ACCENT)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.left_indent = Cm(1.0)
    r3 = p2.add_run(desc)
    set_run_font(r3, size=10, color=INK_SOFT)

para("핵심 기능 우선순위 (v0.1 MVP)", size=12, bold=True, color=ACCENT2, space_before=6, space_after=6)
add_table(
    ["우선순위", "기능", "설명"],
    [
        ["P0", "얼리 오픈 운영", "경쟁 동네 카페보다 이른 개점, 출근 시간대 선점"],
        ["P0", "빠른 픽업 동선", "테이크아웃 전용 픽업존·간편 결제로 대기 최소화"],
        ["P0", "시그니처 모닝 메뉴", "갓 내린 커피 + 간단 아침거리 세트(가격 합리화)"],
        ["P1", "단골 간편 적립·주문 저장", "쿠폰/간단 멤버십으로 단골 인식"],
        ["P1", "사전주문(모바일/전화)", "픽업 시간 지정, 줄 서지 않는 경험"],
        ["P2", "동네 관계 콘텐츠", "원두 스토리·단골 이름 기억 등 비가격 차별화"],
    ],
    widths=[2.0, 5.0, 9.0],
    first_col_bold=True,
)

# =========================================================
# 4. 연락처 및 다음 단계
# =========================================================
heading(4, "연락처 및 다음 단계")

para("다음 단계 (Next Steps)", size=12, bold=True, color=ACCENT2, space_after=6)
nexts = [
    "출근 시간대 매장 앞 동선·대기 현장 관찰 (위치-동선 일치 검증)",
    "단골 고객 5~10명 짧은 인터뷰 (페인·푸드 동반구매 의향 확인)",
    "인근 경쟁 카페 오픈시간·메뉴 가격·픽업 동선 현장 조사",
    "얼리 오픈 인건비 대비 아침 매출 수익성 시범 운영 측정",
    "사전주문·간편 적립 도입 비용 및 운영 부담 검토",
]
for n in nexts:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.6)
    r1 = p.add_run("□ ")
    set_run_font(r1, size=10.5, bold=True, color=ACCENT2)
    r2 = p.add_run(n)
    set_run_font(r2, size=10, color=INK)

para("성공 지표 (Success Metric)", size=12, bold=True, color=ACCENT2, space_before=10, space_after=6)
para("북극성 지표 — 주중 오전(7~9시) 객수 / 재방문 단골 비율. "
     "보조 지표 — 객단가, 테이크아웃 평균 대기 시간, 사전주문 비중, 단골 재방문 주기. "
     "(목표치는 실제 데이터 확보 후 설정)", size=10, color=INK_SOFT, space_after=12)

para("연락처", size=12, bold=True, color=ACCENT2, space_before=4, space_after=6)
add_table(
    ["항목", "내용"],
    [
        ["사업명", "모닝브루 (Morning Brew)"],
        ["위치", "서울 망원동(마포구) 골목 ※ 실습용 예시"],
        ["이메일", "admin@inno-curve.com"],
        ["연락처", "02-000-0000 (실습용 예시 번호)"],
        ["인스타그램", "@morningbrew_demo"],
    ],
    widths=[3.5, 12.5],
    first_col_bold=True,
)

para("", space_after=4)
para("※ 본 문서의 가격·주소·연락처·고객 니즈 서술 중 ‘확인 필요’ 항목은 실제 인터뷰·관찰·매장 데이터로 검증이 필요한 실습용 초안입니다.",
     size=8.5, color=INK_SOFT, align=WD_ALIGN_PARAGRAPH.LEFT)

out = r"C:\Users\samsung-user\Desktop\ttt\Day07\Day_07_사업소개서.docx"
doc.save(out)
print("SAVED:", out)
